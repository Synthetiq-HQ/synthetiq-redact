"""
Document reconstruction and export helpers.

Creates product-facing text and DOCX artifacts from OCR/transcription/redaction
results without making DOCX the source of truth.
"""

from __future__ import annotations

import html
import os
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


def write_text_artifacts(
    out_folder: str,
    doc_id: int,
    raw_text: str,
    clean_text: str,
    redacted_text: str,
) -> dict[str, str]:
    """Write raw, clean, and redacted text artifacts."""
    os.makedirs(out_folder, exist_ok=True)
    paths = {
        "ocr_raw_path": os.path.join(out_folder, "ocr_raw.txt"),
        "transcription_clean_path": os.path.join(out_folder, "transcription_clean.txt"),
        "text_export_path": os.path.join(out_folder, "redacted_text.txt"),
    }
    Path(paths["ocr_raw_path"]).write_text(raw_text or "", encoding="utf-8")
    Path(paths["transcription_clean_path"]).write_text(clean_text or "", encoding="utf-8")
    Path(paths["text_export_path"]).write_text(redacted_text or "", encoding="utf-8")

    # Compatibility aliases for older screens/scripts.
    Path(os.path.join(out_folder, f"{doc_id}_ocr_raw.txt")).write_text(raw_text or "", encoding="utf-8")
    Path(os.path.join(out_folder, f"{doc_id}_export.txt")).write_text(redacted_text or "", encoding="utf-8")
    return paths


def write_transcription_json(out_folder: str, transcription: dict[str, Any]) -> str:
    """Write structured transcription JSON."""
    import json

    path = os.path.join(out_folder, "transcription.json")
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(transcription, handle, indent=2, ensure_ascii=False)
    return path


def write_metadata_json(out_folder: str, metadata: dict[str, Any]) -> str:
    """Write metadata JSON with stable formatting."""
    import json

    path = os.path.join(out_folder, "metadata.json")
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(metadata, handle, indent=2, ensure_ascii=False)
    return path


def write_redacted_docx(
    out_folder: str,
    doc_id: int,
    filename: str,
    clean_text: str,
    redacted_text: str,
    metadata: dict[str, Any],
) -> str:
    """
    Write a staff-friendly DOCX export.

    Uses python-docx if installed; otherwise writes a minimal standards-compliant
    DOCX package directly.
    """
    path = os.path.join(out_folder, "redacted_document.docx")
    try:
        _write_docx_python_docx(path, doc_id, filename, clean_text, redacted_text, metadata)
    except Exception:
        _write_docx_minimal(path, doc_id, filename, clean_text, redacted_text, metadata)
    return path


def _write_docx_python_docx(
    path: str,
    doc_id: int,
    filename: str,
    clean_text: str,
    redacted_text: str,
    metadata: dict[str, Any],
) -> None:
    """Write DOCX using python-docx when available."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Hillingdon Document Processor", level=1)
    doc.add_paragraph(f"Document ID: {doc_id}")
    doc.add_paragraph(f"Original filename: {filename}")
    doc.add_paragraph(f"Category: {metadata.get('category') or 'unknown'}")
    doc.add_paragraph(f"Redaction profile: {metadata.get('redaction_profile') or '-'}")
    doc.add_paragraph(f"Review required: {metadata.get('flag_needs_review')}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
    doc.add_heading("Redacted Transcription", level=2)
    for line in (redacted_text or "").splitlines() or [""]:
        doc.add_paragraph(line)
    doc.add_heading("Clean Transcription", level=2)
    for line in (clean_text or "").splitlines() or [""]:
        doc.add_paragraph(line)
    doc.save(path)


def _write_docx_minimal(
    path: str,
    doc_id: int,
    filename: str,
    clean_text: str,
    redacted_text: str,
    metadata: dict[str, Any],
) -> None:
    """Write a minimal DOCX zip package without external dependencies."""
    paragraphs = [
        "Hillingdon Document Processor",
        f"Document ID: {doc_id}",
        f"Original filename: {filename}",
        f"Category: {metadata.get('category') or 'unknown'}",
        f"Redaction profile: {metadata.get('redaction_profile') or '-'}",
        f"Review required: {metadata.get('flag_needs_review')}",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        "",
        "Redacted Transcription",
        *(redacted_text or "").splitlines(),
        "",
        "Clean Transcription",
        *(clean_text or "").splitlines(),
    ]
    document_xml = _minimal_document_xml(paragraphs)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)
        docx.writestr("_rels/.rels", _RELS_XML)
        docx.writestr("word/document.xml", document_xml)


def _minimal_document_xml(paragraphs: list[str]) -> str:
    """Build Word document XML for simple paragraphs."""
    body = "\n".join(
        f"<w:p><w:r><w:t xml:space=\"preserve\">{html.escape(paragraph)}</w:t></w:r></w:p>"
        for paragraph in paragraphs
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>
    </w:sectPr>
  </w:body>
</w:document>"""


_CONTENT_TYPES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

_RELS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
